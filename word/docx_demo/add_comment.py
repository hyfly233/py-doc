import os
import re
from typing import List

from docx import Document
from dotenv import load_dotenv

load_dotenv()


def copy_font_format(source_font, target_font):
    """复制字体格式"""
    try:
        if source_font.name:
            target_font.name = source_font.name
        if source_font.size:
            target_font.size = source_font.size
        if source_font.bold is not None:
            target_font.bold = source_font.bold
        if source_font.italic is not None:
            target_font.italic = source_font.italic
        if source_font.underline is not None:
            target_font.underline = source_font.underline
        if source_font.color.rgb:
            target_font.color.rgb = source_font.color.rgb
    except:
        pass


def add_comments(file_path: str, words: List[str]):
    if not file_path.endswith('.docx'):
        print("只支持 .docx 格式文件")
        return
    doc = Document(file_path)

    # 生成新文件名
    base_name = os.path.splitext(file_path)[0]
    new_file_path = f"{base_name}_comments.docx"

    if words is None or words is False or len(words) == 0:
        print("没有提供任何词语进行添加注释")
        return

    # 按词语长度降序排序，优先处理长词，避免短词干扰长词的匹配
    sorted_words = sorted(words, key=len, reverse=True)

    for i, paragraph in enumerate(doc.paragraphs):
        full_text = paragraph.text

        # 检查是否包含任何目标词语
        has_target_words = any(word in full_text for word in sorted_words)
        if not has_target_words:
            continue

        print(f"index: {i}, original text: {full_text}")

        # 创建字符到格式的映射
        char_formats = []
        char_index = 0

        for run in paragraph.runs:
            run_text = run.text
            for char in run_text:
                char_formats.append(run.font)
                char_index += 1

        # 构建正则表达式，匹配所有目标词语
        pattern = '|'.join(re.escape(word) for word in sorted_words)

        # 分割文本，保留分隔符
        parts = re.split(f'({pattern})', full_text)

        # 清空段落的所有runs
        for run in paragraph.runs:
            run.clear()

        # 移除所有runs
        while len(paragraph.runs) > 0:
            paragraph._element.remove(paragraph.runs[0]._element)

        # 重新构建段落内容
        current_pos = 0
        for part in parts:
            if not part:  # 跳过空字符串
                continue

            if part in sorted_words:
                # 添加目标词（需要添加注释的词）
                target_run = paragraph.add_run(part)

                # 继承原文字的格式
                if current_pos < len(char_formats):
                    source_font = char_formats[current_pos]
                    copy_font_format(source_font, target_run.font)

                # 为目标词添加注释
                try:
                    doc.add_comment(
                        runs=[target_run],
                        text=f"这是对词语'{part}'的评论。",
                        author="作者名(测试)",
                        initials="作者(测试)"
                    )
                except Exception as e:
                    print(f"添加注释失败: {e}")
                    # 如果添加注释失败，至少保持文本不变
            else:
                # 对于普通文本，按字符重建，保持原有格式
                start_pos = current_pos
                for char_idx, char in enumerate(part):
                    format_idx = start_pos + char_idx
                    if format_idx < len(char_formats):
                        # 为每个字符创建单独的run（如果格式不同）
                        if char_idx == 0 or (
                                format_idx > 0 and char_formats[format_idx] != char_formats[format_idx - 1]):
                            # 创建新的run
                            char_run = paragraph.add_run(char)
                            copy_font_format(char_formats[format_idx], char_run.font)
                        else:
                            # 添加到最后一个run
                            if paragraph.runs:
                                paragraph.runs[-1].text += char
                    else:
                        # 如果超出了格式范围，使用最后一个可用格式
                        if char_idx == 0:
                            char_run = paragraph.add_run(char)
                            if char_formats:
                                copy_font_format(char_formats[-1], char_run.font)
                        else:
                            if paragraph.runs:
                                paragraph.runs[-1].text += char

            current_pos += len(part)

        print(f"index: {i}, processed text: {paragraph.text}")

    # 保存新文档
    doc.save(new_file_path)
    print(f"标注完成，新文件已保存为: {new_file_path}")


if __name__ == '__main__':
    word_path: str = os.getenv('WORD_PATH')
    add_comments(word_path, ["喵", "公司", "北京", "喵喵公司"])
