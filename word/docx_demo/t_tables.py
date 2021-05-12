import os

from docx import Document
from dotenv import load_dotenv

load_dotenv()


def table_to_markdown(table):
    """将 python-docx 表格对象转换为 Markdown 字符串"""
    markdown_lines = []

    # 处理表格行
    for i, row in enumerate(table.rows):
        # 提取单元格文本
        cells = []
        for cell in row.cells:
            # 获取单元格文本，移除多余空白
            cell_text = cell.text.strip().replace('\n', ' ')
            cells.append(cell_text)

        # 创建表格行
        markdown_line = '| ' + ' | '.join(cells) + ' |'
        markdown_lines.append(markdown_line)

        # # 在第一行后添加分隔符
        # if i == 0:
        #     separator = '|' + '|'.join([' --- ' for _ in cells]) + '|'
        #     markdown_lines.append(separator)

    return '\n'.join(markdown_lines)


def convert_docx_tables_to_markdown(docx_path):
    """读取 Word 文档中的所有表格并转换为 Markdown"""
    doc = Document(docx_path)
    markdown_content = []

    for i, table in enumerate(doc.tables):
        markdown_content.append(f"## 表格 {i + 1}")
        markdown_content.append(table_to_markdown(table))
        markdown_content.append("")  # 添加空行

    return '\n'.join(markdown_content)


def parse_document_tables(file_path):
    if file_path.endswith('.docx'):
        doc = Document(file_path)

        print(f"表格个数: {len(doc.tables)}")

        for i, table in enumerate(doc.tables):
            for j, row in enumerate(table.rows):
                for k, cell in enumerate(row.cells):
                    print(f"表格 {i + 1}，行 {j + 1}，列 {k + 1}，内容: {cell.text}")

                print("----------")

            print("=====================")

        return doc


# 使用示例
if __name__ == "__main__":
    word_path: str = os.getenv('WORD_PATH')
    parse_document_tables(word_path)

    print("\n----------------------------------------------\n")

    markdown_result = convert_docx_tables_to_markdown(word_path)
    print(markdown_result)

    # 保存到文件
    # with open("output.md", "w", encoding="utf-8") as f:
    #     f.write(markdown_result)
