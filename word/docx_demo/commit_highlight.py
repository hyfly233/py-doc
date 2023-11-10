import os
import re
from dataclasses import dataclass
from typing import Optional, Tuple, List, Dict

from docx import Document
from docx.text.paragraph import Paragraph
from dotenv import load_dotenv

from word.docx_demo.common import copy_font_format

load_dotenv()


@dataclass
class AnnotationConfig:
    """标注配置类"""

    add_comment: bool = False  # 是否添加注释
    comment_text: str = ""  # 注释内容
    comment_author: str = "标注者"  # 注释作者
    comment_initials: str = "标"  # 作者简称

    highlight: bool = False  # 是否高亮
    highlight_color: str = "yellow"  # 高亮颜色名称

    emphasize: bool = False  # 是否突出显示 (添加括号)
    emphasize_symbols: Tuple[str, str] = ("「", "」")  # 突出显示符号

    font_color: Optional[str] = None  # 字体颜色名称


def _add_comment(doc, run, config: AnnotationConfig, word: str):
    """
    添加注释

    Args:
        doc: 当前文档对象
        run: 目标run
        config: 标注配置
        word: 被注释的词语

    Returns:
        None
    """
    try:
        comment_text = config.comment_text or f"标注词语: {word}"
        doc.add_comment(
            runs=[run],
            text=comment_text,
            author=config.comment_author,
            initials=config.comment_initials,
        )
    except Exception as e:
        print(f"添加注释失败: {e}")


def _add_normal_text(
    paragraph: Paragraph, text: str, char_formats: List, start_pos: int
):
    """
    添加普通文本，保持原有格式

    Args:
        paragraph: 目标段落
        text: 目标文本
        char_formats: 每个字符对应的格式列表
        start_pos: 当前文本在段落中的起始位置

    Returns:
        None
    """
    for char_idx, char in enumerate(text):
        format_idx = start_pos + char_idx

        if format_idx < len(char_formats):
            # 检查是否需要创建新run
            need_new_run = char_idx == 0 or (
                format_idx > 0
                and char_formats[format_idx] != char_formats[format_idx - 1]
            )

            if need_new_run:
                char_run = paragraph.add_run(char)
                copy_font_format(char_formats[format_idx], char_run.font)
                # 保持原有颜色，copy_font_format 没有复制颜色
                if char_formats[format_idx].color.rgb:
                    char_run.font.color.rgb = char_formats[format_idx].color.rgb
            else:
                # 添加到最后一个run
                if paragraph.runs:
                    paragraph.runs[-1].text += char
        else:
            # 使用最后可用格式
            _add_char_with_fallback_format(paragraph, char, char_formats, char_idx)


def _add_char_with_fallback_format(
    paragraph: Paragraph, char: str, char_formats: List, char_idx: int
):
    """
    使用备用格式添加字符

    Args:
        paragraph: 目标段落
        char: 目标字符
        char_formats: 每个字符对应的格式列表
        char_idx: 当前字符索引

    Returns:
        None
    """
    if char_idx == 0:
        char_run = paragraph.add_run(char)
        if char_formats:
            copy_font_format(char_formats[-1], char_run.font)
            if char_formats[-1].color.rgb:
                char_run.font.color.rgb = char_formats[-1].color.rgb
    else:
        if paragraph.runs:
            paragraph.runs[-1].text += char


def _apply_annotation_format(run, config: AnnotationConfig):
    """
    应用标注格式

    Args:
        run: 目标run
        config: 标注配置

    Returns:
        None
    """
    # 修改字体颜色
    if config.font_color:
        _apply_font_color(run, config.font_color)

    # 高亮
    if config.highlight:
        _apply_highlight_color(run, config.highlight_color)


def _apply_highlight_color(run, color_name: str):
    """
    应用高亮颜色（使用颜色名称）

    Args:
        run: 目标run
        color_name: 颜色名称

    Returns:
        None
    """
    from docx.enum.text import WD_COLOR_INDEX

    color_map = {
        "yellow": WD_COLOR_INDEX.YELLOW,
        "red": WD_COLOR_INDEX.RED,
        "green": WD_COLOR_INDEX.BRIGHT_GREEN,
        "blue": WD_COLOR_INDEX.BLUE,
        "pink": WD_COLOR_INDEX.PINK,
        "cyan": WD_COLOR_INDEX.TURQUOISE,
        "gray": WD_COLOR_INDEX.GRAY_25,
        "purple": WD_COLOR_INDEX.VIOLET,
        "lime": WD_COLOR_INDEX.BRIGHT_GREEN,
    }

    run.font.highlight_color = color_map.get(color_name.lower(), WD_COLOR_INDEX.YELLOW)


def _apply_font_color(run, color_name: str):
    """
    应用字体颜色（使用颜色名称）

    Args:
        run: 目标run
        color_name: 颜色名称

    Returns:
        None
    """
    from docx.shared import RGBColor

    color_map = {
        "red": RGBColor(255, 0, 0),
        "blue": RGBColor(0, 0, 255),
        "green": RGBColor(0, 128, 0),
        "purple": RGBColor(128, 0, 128),
        "brown": RGBColor(165, 42, 42),
        "black": RGBColor(0, 0, 0),
        "gray": RGBColor(128, 128, 128),
        "pink": RGBColor(255, 192, 203),
        "yellow": RGBColor(255, 255, 0),
    }

    run.font.color.rgb = color_map.get(color_name.lower(), RGBColor(0, 0, 0))


def _create_char_format_mapping(paragraph: Paragraph) -> List:
    """
    创建字符到格式的映射

    Args:
        paragraph: 目标段落

    Returns:
        List 每个字符对应的格式列表
    """
    char_formats = []
    for run in paragraph.runs:
        for _ in run.text:
            char_formats.append(run.font)
    return char_formats


def _clear_paragraph_runs(paragraph: Paragraph):
    """
    清空段落的所有runs

    Args:
        paragraph: 目标段落

    Returns:
        None
    """
    for run in paragraph.runs:
        run.clear()

    while len(paragraph.runs) > 0:
        paragraph._element.remove(paragraph.runs[0]._element)

    # 从后往前删除，避免索引问题
    # for i in range(len(paragraph.runs) - 1, -1, -1):
    #     paragraph.runs[i].clear()
    #     # 使用公共方法删除run
    #     paragraph._p.remove(paragraph.runs[i]._r)


class DocumentAnnotator:
    """文档标注器"""

    def __init__(self, word_configs: Dict[str, AnnotationConfig]):
        self.word_configs = word_configs
        # 将词语按长度排序，防止子串冲突，如 "喵喵公司" "公司" "喵"
        self.sorted_words = sorted(word_configs.keys(), key=len, reverse=True)
        # 构建正则表达式模式，如 "喵喵公司|公司|喵"
        self.pattern = "|".join(re.escape(word) for word in self.sorted_words)

    def annotate_document(self, file_path: str) -> str:
        """
        标注整个文档

        Args:
            file_path: 源文档路径

        Returns:
            None
        """
        if not file_path.endswith(".docx"):
            raise ValueError("只支持 .docx 格式文件")

        # 生成新文件名
        base_name = os.path.splitext(file_path)[0]
        new_file_path = f"{base_name}_标注版本.docx"

        try:
            doc = Document(file_path)

            # 处理段落
            print(f"段落数量: {len(doc.paragraphs)}")
            self._process_paragraphs(doc.paragraphs, doc)

            # 处理表格
            print(f"表格数量: {len(doc.tables)}")
            self._process_tables(doc.tables, doc)

            # 保存文档
            doc.save(new_file_path)
            print(f"标注完成，新文件已保存为: {new_file_path}")
            return new_file_path

        except Exception as e:
            print(f"处理文档时发生错误: {e}")
            raise

    def _process_tables(self, tables, doc):
        """
        处理所有表格

        Args:
            tables: 文档中的表格列表
            doc: 当前文档对象

        Returns:
            None
        """
        for i, table in enumerate(tables):
            for j, row in enumerate(table.rows):
                for k, cell in enumerate(row.cells):
                    if cell.text.strip():  # 只处理非空单元格
                        print(
                            f"表格 {i + 1}, 行 {j + 1}, 列 {k + 1}, 内容: {cell.text}"
                        )
                        self._process_paragraphs(cell.paragraphs, doc)

    def _process_paragraphs(self, paragraphs: List[Paragraph], doc):
        """
        处理段落列表

        Args:
            paragraphs: 目标段落列表
            doc: 当前文档对象

        Returns:
            None
        """
        for i, paragraph in enumerate(paragraphs):
            self._process_single_paragraph(paragraph, doc, i)

    def _process_single_paragraph(self, paragraph: Paragraph, doc, index: int):
        """
        处理单个段落

        Args:
            paragraph: 目标段落

        Returns:
            None
        """
        full_text = paragraph.text

        # 检查是否包含任何目标词语
        if not any(word in full_text for word in self.sorted_words):
            return

        print(f"段落 {index + 1}, 内容: {full_text}")

        try:
            # 创建字符到格式的映射
            char_formats = _create_char_format_mapping(paragraph)

            # 分割文本
            parts = re.split(f"({self.pattern})", full_text)

            # 清空并重建段落
            _clear_paragraph_runs(paragraph)
            self._rebuild_paragraph(paragraph, parts, char_formats, doc)

            print(f"  段落 {index + 1}, 处理后的内容: {paragraph.text}")

        except Exception as e:
            print(f"处理段落 ({index + 1})时发生错误: {e}")

    def _rebuild_paragraph(
        self, paragraph: Paragraph, parts: List[str], char_formats: List, doc
    ):
        """
        重建段落内容

        Args:
            paragraph: 目标段落
            parts: 分割后的文本部分列表
            char_formats: 每个字符对应的格式列表
            doc: 当前文档对象

        Returns:
            None
        """
        current_pos = 0

        for part in parts:
            if not part:  # 跳过空字符串
                continue

            if part in self.sorted_words:
                self._add_annotated_word(
                    paragraph, part, char_formats, current_pos, doc
                )
            else:
                _add_normal_text(paragraph, part, char_formats, current_pos)

            current_pos += len(part)

    def _add_annotated_word(
        self, paragraph: Paragraph, word: str, char_formats: List, current_pos: int, doc
    ):
        """
        添加标注的词语

        Args:
            paragraph: 目标段落
            word: 目标词语
            char_formats: 每个字符对应的格式列表
            current_pos: 当前字符位置
            doc: 当前文档对象

        Returns:
            None
        """
        config = self.word_configs[word]

        # 构建标注文本
        annotated_word = word

        # 突出显示 (添加括号)
        if config.emphasize:
            annotated_word = (
                f"{config.emphasize_symbols[0]}{word}{config.emphasize_symbols[1]}"
            )

        # 添加run
        target_run = paragraph.add_run(annotated_word)

        # 继承原格式
        if current_pos < len(char_formats):
            # 找出当前词语原有的格式，复制到新 run
            copy_font_format(char_formats[current_pos], target_run.font)

        # 应用标注格式
        _apply_annotation_format(target_run, config)

        # 添加注释
        if config.add_comment:
            _add_comment(doc, target_run, config, word)


def annotate_words_with_configs(
    file_path: str, word_configs: Dict[str, AnnotationConfig]
) -> str:
    """
    在文档中标注多个词语（为每个词语使用不同配置）

    Args:
        file_path: 源文档路径
        word_configs: 词语和配置的字典 {词语: 配置}

    Returns:
        新文件路径
    """
    if not word_configs:
        raise ValueError("没有提供任何词语进行标注")

    annotator = DocumentAnnotator(word_configs)
    return annotator.annotate_document(file_path)


def annotate_multiple_words_same_config(
    file_path: str, target_words: List[str], config: AnnotationConfig
) -> str:
    """
    在文档中标注多个词语（使用相同配置）

    Args:
        file_path: 源文档路径
        target_words: 要标注的词语列表
        config: 标注配置

    Returns:
        新文件路径
    """
    word_configs = {word: config for word in target_words}
    return annotate_words_with_configs(file_path, word_configs)


if __name__ == "__main__":
    word_path: str = os.getenv("WORD_PATH")

    # 方案1: 多个词语使用不同配置
    word_configs = {
        "附件": AnnotationConfig(highlight=True, highlight_color="red"),
        "公司": AnnotationConfig(highlight=True, highlight_color="yellow"),
        "喵": AnnotationConfig(font_color="red"),
        "卖方": AnnotationConfig(
            add_comment=True,
            comment_text="对卖方的评论",
            comment_author="评论员1",
            highlight=True,
            highlight_color="green",
            font_color="blue",
        ),
        "喵喵公司": AnnotationConfig(
            add_comment=True,
            comment_text="对喵喵公司的评论",
            comment_author="评论员2",
            highlight=True,
            highlight_color="red",
            font_color="lime",
        ),
    }

    try:
        new_file = annotate_words_with_configs(word_path, word_configs)
        print(f"处理完成，生成文件: {new_file}")
    except Exception as e:
        print(f"处理失败: {e}")

    # 方案2: 多个词语使用相同配置
    # try:
    #     same_config = create_emphasize_config(color="red")
    #     new_file = annotate_multiple_words_same_config(word_path, ["甲方", "乙方"], same_config)
    #     print(f"处理完成，生成文件: {new_file}")
    # except Exception as e:
    #     print(f"处理失败: {e}")
