import os

from docx import Document
from dotenv import load_dotenv

load_dotenv()


def parse_document_paragraphs(file_path):
    if file_path.endswith(".docx"):
        doc = Document(file_path)

        print(f"段落数量: {len(doc.paragraphs)}")

        for i, para in enumerate(doc.paragraphs):
            print(f"段落 {i + 1}，内容: {para.text}")

            for j, run in enumerate(para.runs):
                print(f"   段落 {i + 1}，Run {j + 1}，内容: {run.text}")

            print("---------------------")


if __name__ == "__main__":
    word_path: str = os.getenv("WORD_PATH")
    parse_document_paragraphs(word_path)
